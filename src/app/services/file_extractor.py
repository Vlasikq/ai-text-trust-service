"""Извлечение plain-text из загруженных файлов: TXT / DOCX / PDF.

Контракт:
  - `extract(file_bytes, content_type, filename)` → `ExtractResult`.
  - Все ошибки парсинга → `ExtractError` с понятным русскоязычным detail.
  - Encrypted PDF / scanned PDF (no text layer) → `ExtractError` с явным сообщением.

Privacy: модуль не пишет в БД и не логирует контент. Извлечённый текст
возвращается клиенту, который сам вставляет его в форму анализа.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Literal

import chardet

SourceFormat = Literal["txt", "docx", "pdf"]


class ExtractError(ValueError):
    """Не удалось извлечь текст: encrypted/scanned PDF, битый DOCX, неизвестный тип."""


@dataclass
class ExtractResult:
    text: str
    source_format: SourceFormat
    char_count: int
    warnings: list[str] = field(default_factory=list)


def _detect_format(content_type: str | None, filename: str | None) -> SourceFormat:
    """Определить формат: сначала по content-type, потом по расширению."""
    ct = (content_type or "").lower()
    name = (filename or "").lower()
    if "wordprocessingml" in ct or name.endswith(".docx"):
        return "docx"
    if "pdf" in ct or name.endswith(".pdf"):
        return "pdf"
    if (
        ct.startswith("text/")
        or name.endswith(".txt")
        or name.endswith(".md")
        or not ct
    ):
        return "txt"
    raise ExtractError(f"Неподдерживаемый тип файла: {content_type or filename}")


def extract(
    data: bytes,
    *,
    content_type: str | None = None,
    filename: str | None = None,
) -> ExtractResult:
    if not data:
        raise ExtractError("Файл пуст")
    fmt = _detect_format(content_type, filename)
    if fmt == "txt":
        return _extract_txt(data)
    if fmt == "docx":
        return _extract_docx(data)
    return _extract_pdf(data)


# Извлечение TXT.


def _extract_txt(data: bytes) -> ExtractResult:
    """UTF-8 с BOM → utf-8-sig; иначе chardet; fallback utf-8 с заменой битых байт."""
    warnings: list[str] = []
    if data.startswith(b"\xef\xbb\xbf"):
        encoding = "utf-8-sig"
    else:
        det = chardet.detect(data) or {}
        encoding = det.get("encoding") or "utf-8"
        if encoding.lower() == "ascii":
            encoding = "utf-8"
    try:
        text = data.decode(encoding, errors="strict")
    except UnicodeDecodeError:
        text = data.decode(encoding, errors="replace")
        warnings.append(
            f"Часть символов не декодировались в {encoding} и были заменены"
        )
    text = text.strip()
    return ExtractResult(
        text=text,
        source_format="txt",
        char_count=len(text),
        warnings=warnings,
    )


# Извлечение DOCX.


def _extract_docx(data: bytes) -> ExtractResult:
    """python-docx: paragraphs + table cells. Изображения/прочее игнорируем."""
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover
        raise ExtractError("python-docx не установлен") from e

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractError(f"Не удалось прочитать DOCX: {e}") from e

    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    parts.append(cell_text)

    text = "\n\n".join(parts).strip()
    if not text:
        raise ExtractError("В DOCX не найдено ни одного абзаца с текстом")
    return ExtractResult(text=text, source_format="docx", char_count=len(text))


# Извлечение PDF.


def _extract_pdf(data: bytes) -> ExtractResult:
    """pypdf: text-layer extraction. Encrypted PDF → ошибка. Scan-PDF (нет
    текстового слоя) — возвращаем пустой text → ExtractError выше уровнем."""
    try:
        import pypdf
    except ImportError as e:  # pragma: no cover
        raise ExtractError("pypdf не установлен") from e

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as e:
        raise ExtractError(f"Не удалось прочитать PDF: {e}") from e

    if reader.is_encrypted:
        # pypdf пытается decrypt пустым паролем — если не удалось, это шифр.
        try:
            reader.decrypt("")
        except Exception:
            pass
        if reader.is_encrypted:
            raise ExtractError(
                "PDF защищён паролем. Снимите защиту и загрузите снова."
            )

    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            parts.append(page_text)

    text = "\n\n".join(parts).strip()
    if not text:
        raise ExtractError(
            "PDF не содержит текстового слоя (вероятно, скан или изображение). "
            "Распознайте OCR'ом и загрузите как TXT/DOCX."
        )
    return ExtractResult(text=text, source_format="pdf", char_count=len(text))
